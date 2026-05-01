#!/usr/bin/env python
"""Generate report.docx from report_draft.md using python-docx.

The output is formatted for Microsoft Word on Windows.
"""

import re
import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

INPUT = os.path.join(os.path.dirname(os.path.abspath(__file__)), "report_draft.md")
OUTPUT = os.path.join(os.path.dirname(os.path.abspath(__file__)), "report.docx")

# Font name for Korean
FONT = "Malgun Gothic"

FIGURES = {
    "개념도. 연구 전체 흐름": {
        "path": os.path.join("results", "report_figures", "figure0_pipeline_overview.png"),
        "width": 6.6,
        "show_caption": False,
    },
    "그림 1. 시간과 자원 효율의 핵심 비교": {
        "path": os.path.join("results", "report_figures", "figure1_time_efficiency_summary.png"),
        "width": 6.5,
    },
    "그림 2. 장애 강도에 따른 미도착 인원 변화": {
        "path": os.path.join("results", "report_figures", "figure2_undelivered_risk.png"),
        "width": 6.3,
    },
    "그림 3. 의사결정 관점 요약": {
        "path": os.path.join("results", "report_figures", "figure3_decision_lens.png"),
        "width": 6.4,
    },
}


def set_run_font(run, size=11, bold=False, color=None):
    """Apply consistent Korean font to a run."""
    run.font.name = FONT
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = color
    # Set East Asian font
    rpr = run._element.get_or_add_rPr()
    rFonts = rpr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = rpr.makeelement(qn("w:rFonts"), {})
        rpr.insert(0, rFonts)
    rFonts.set(qn("w:eastAsia"), FONT)
    rFonts.set(qn("w:ascii"), FONT)
    rFonts.set(qn("w:hAnsi"), FONT)


def add_styled_paragraph(doc, text, style=None, font_size=11, bold=False, color=None):
    """Add paragraph with Korean font, handling **bold** inline."""
    p = doc.add_paragraph(style=style)
    # Parse bold segments
    parts = re.split(r"(\*\*.+?\*\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = p.add_run(part[2:-2])
            set_run_font(run, size=font_size, bold=True, color=color)
        elif part:
            run = p.add_run(part)
            set_run_font(run, size=font_size, bold=bold, color=color)
    return p


def add_table(doc, rows):
    """Add a formatted table."""
    if not rows:
        return
    ncols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=ncols)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, row in enumerate(rows):
        for j, cell_text in enumerate(row):
            cell = table.cell(i, j)
            cell.text = ""
            p = cell.paragraphs[0]
            text = cell_text.strip()
            run = p.add_run(text)
            is_header = (i == 0)
            set_run_font(run, size=10, bold=is_header,
                         color=RGBColor(0xFF, 0xFF, 0xFF) if is_header else None)
            if is_header:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                # Header cell shading
                shading = cell._element.get_or_add_tcPr()
                shd = shading.makeelement(qn("w:shd"), {
                    qn("w:val"): "clear",
                    qn("w:color"): "auto",
                    qn("w:fill"): "1F3864"
                })
                shading.append(shd)
    # Add spacing after table
    doc.add_paragraph()


def add_figure(doc, caption, figure_spec):
    """Add a report figure with a centered caption."""
    base_dir = os.path.dirname(os.path.abspath(__file__))
    image_path = os.path.join(base_dir, figure_spec["path"])
    if not os.path.exists(image_path):
        p = add_styled_paragraph(doc, caption, font_size=9, color=RGBColor(0x66, 0x66, 0x66))
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        return

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(image_path, width=Inches(figure_spec["width"]))

    if figure_spec.get("show_caption", True):
        caption_p = add_styled_paragraph(
            doc,
            caption,
            font_size=9,
            bold=False,
            color=RGBColor(0x66, 0x66, 0x66),
        )
        caption_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()


def parse_and_generate(md_path, docx_path):
    """Parse markdown and generate docx."""
    with open(md_path, "r", encoding="utf-8") as f:
        lines = f.readlines()

    doc = Document()

    # Page setup: A4, 2.54cm margins
    for section in doc.sections:
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)

    # Default font
    style = doc.styles["Normal"]
    style.font.name = FONT
    style.font.size = Pt(11)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), FONT)

    i = 0
    while i < len(lines):
        line = lines[i].rstrip("\n")
        stripped = line.strip()

        # Blank line
        if not stripped:
            i += 1
            continue

        # Figure captions mapped to project-local report images.
        if stripped in FIGURES:
            add_figure(doc, stripped, FIGURES[stripped])
            i += 1
            continue

        # Horizontal rule
        if stripped == "---":
            # Add a thin line paragraph
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run("-" * 60)
            set_run_font(run, size=8, color=RGBColor(0xCC, 0xCC, 0xCC))
            i += 1
            continue

        # Heading
        hmatch = re.match(r"^(#{1,6})\s+(.+)$", stripped)
        if hmatch:
            level = len(hmatch.group(1))
            text = hmatch.group(2)
            if level == 1:
                add_styled_paragraph(doc, text, font_size=18, bold=True)
            elif level == 2:
                add_styled_paragraph(doc, text, font_size=14, bold=True)
            elif level == 3:
                add_styled_paragraph(doc, text, font_size=12, bold=True)
            else:
                add_styled_paragraph(doc, text, font_size=11, bold=True)
            i += 1
            continue

        # Table
        if stripped.startswith("|"):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i].strip())
                i += 1
            rows = []
            for tl in table_lines:
                cells = [c.strip() for c in tl.split("|")[1:-1]]
                if all(re.match(r"^[-:]+$", c) for c in cells):
                    continue
                rows.append(cells)
            if rows:
                add_table(doc, rows)
            continue

        # Bullet point
        if re.match(r"^[-*+]\s+", stripped):
            text = re.sub(r"^[-*+]\s+", "", stripped)
            # Use List Bullet style
            add_styled_paragraph(doc, text, style="List Bullet", font_size=11)
            i += 1
            continue

        # Blockquote (>)
        if stripped.startswith(">"):
            text = stripped.lstrip("> ").strip()
            p = add_styled_paragraph(doc, text, font_size=11)
            p.paragraph_format.left_indent = Cm(1.0)
            i += 1
            continue

        # Regular paragraph - collect consecutive non-special lines
        para_lines = []
        while i < len(lines):
            ln = lines[i].rstrip("\n").strip()
            if not ln or ln.startswith("#") or ln.startswith("|") or ln.startswith("---") or re.match(r"^[-*+]\s+", ln) or ln.startswith(">"):
                break
            para_lines.append(ln)
            i += 1
        if para_lines:
            text = " ".join(para_lines)
            add_styled_paragraph(doc, text, font_size=11)

    doc.save(docx_path)
    print(f"Generated: {docx_path} ({os.path.getsize(docx_path):,} bytes)")
if __name__ == "__main__":
    parse_and_generate(INPUT, OUTPUT)
